from __future__ import annotations

from pathlib import Path

import pytest
from brain_core.ingest.handlers.transcript_docx import TranscriptDOCXHandler
from brain_core.ingest.types import SourceType
from docx import Document


# ---------------------------------------------------------------------------
# Fixture builders (Plan 24 T1.5: content-sniff-aware)
# ---------------------------------------------------------------------------


def _build_speaker_transcript(tmp_path: Path, name: str = "speakers.docx") -> Path:
    """5 paragraphs of "Speaker N: text" → expected to satisfy the sniff."""
    doc = Document()
    doc.add_paragraph("Speaker 1: Welcome everyone to today's session.")
    doc.add_paragraph("Speaker 2: Thanks for having us.")
    doc.add_paragraph("Speaker 1: Let's begin with introductions.")
    doc.add_paragraph("Speaker 2: I'm Bob from engineering.")
    doc.add_paragraph("Speaker 1: Great, and I'm Alice from product.")
    path = tmp_path / name
    doc.save(str(path))
    return path


def _build_timestamp_transcript(tmp_path: Path, name: str = "stamps.docx") -> Path:
    """5 paragraphs starting with [HH:MM] timestamps → expected to satisfy the sniff."""
    doc = Document()
    doc.add_paragraph("[00:00] Opening remarks from the chair.")
    doc.add_paragraph("[01:23] First question from the audience.")
    doc.add_paragraph("[02:45] Response from the speaker.")
    doc.add_paragraph("[04:10] Follow-up discussion.")
    doc.add_paragraph("[05:30] Closing remarks.")
    path = tmp_path / name
    doc.save(str(path))
    return path


def _build_generic_docx(tmp_path: Path, name: str = "generic.docx") -> Path:
    """Plain prose paragraphs — no speaker labels, no timestamps."""
    doc = Document()
    doc.add_paragraph("This is the introduction to our quarterly report.")
    doc.add_paragraph("Revenue grew by twelve percent year over year.")
    doc.add_paragraph("We expanded into three new geographic markets.")
    doc.add_paragraph("The team grew by six new hires across all departments.")
    doc.add_paragraph("Looking ahead, we expect continued steady growth.")
    path = tmp_path / name
    doc.save(str(path))
    return path


def _build_short_docx(tmp_path: Path, name: str = "short.docx") -> Path:
    """Only 2 non-empty paragraphs — below the 3-paragraph minimum."""
    doc = Document()
    doc.add_paragraph("Alice: Hello.")
    doc.add_paragraph("Bob: Hi.")
    path = tmp_path / name
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Existing behavior — still works with the new sniffer
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_docx_transcript_reads_paragraphs(fixtures_dir: Path, tmp_path: Path) -> None:
    h = TranscriptDOCXHandler()
    assert h.can_handle(fixtures_dir / "notes.docx")
    es = await h.extract(fixtures_dir / "notes.docx", archive_root=tmp_path)
    assert es.source_type is SourceType.TRANSCRIPT
    assert "Meeting notes 2026-04-13" in es.body_text
    assert "Alice: Welcome to the meeting." in es.body_text
    assert "Bob: Thanks for setting this up." in es.body_text
    assert es.archive_path.exists()


def test_docx_handler_rejects_txt(tmp_path: Path) -> None:
    f = tmp_path / "x.txt"
    f.write_text("nope", encoding="utf-8")
    assert TranscriptDOCXHandler().can_handle(f) is False


@pytest.mark.asyncio
async def test_docx_handler_raises_handler_error_on_corrupt_docx(tmp_path: Path) -> None:
    """A file with .docx extension that is not a valid DOCX must raise HandlerError."""
    from brain_core.ingest.handlers.base import HandlerError

    fake = tmp_path / "fake.docx"
    fake.write_bytes(b"not a docx")
    with pytest.raises(HandlerError):
        await TranscriptDOCXHandler().extract(fake, archive_root=tmp_path / "archive")


# ---------------------------------------------------------------------------
# Plan 24 T1.5 — content-sniffing can_handle
# ---------------------------------------------------------------------------


def test_can_handle_returns_true_for_transcript_shape(tmp_path: Path) -> None:
    """5 paragraphs of 'Speaker N: text' should satisfy the content sniff."""
    path = _build_speaker_transcript(tmp_path)
    assert TranscriptDOCXHandler().can_handle(path) is True


def test_can_handle_returns_true_for_timestamp_shape(tmp_path: Path) -> None:
    """5 paragraphs starting with [HH:MM] timestamps should satisfy the sniff."""
    path = _build_timestamp_transcript(tmp_path)
    assert TranscriptDOCXHandler().can_handle(path) is True


def test_can_handle_returns_false_for_generic_docx(tmp_path: Path) -> None:
    """Plain prose with no speakers / timestamps falls through to DocxHandler."""
    path = _build_generic_docx(tmp_path)
    assert TranscriptDOCXHandler().can_handle(path) is False


def test_can_handle_returns_false_for_short_docx(tmp_path: Path) -> None:
    """Below the 3-paragraph minimum, can_handle returns False regardless of shape."""
    path = _build_short_docx(tmp_path)
    assert TranscriptDOCXHandler().can_handle(path) is False


def test_can_handle_returns_false_for_corrupt_docx(tmp_path: Path) -> None:
    """Garbage bytes with a .docx extension must NOT crash the dispatcher.

    can_handle returns False; the dispatcher then tries the next handler
    (DocxHandler), which surfaces a HandlerError at extract time.
    """
    fake = tmp_path / "corrupt.docx"
    fake.write_bytes(b"not a real docx file at all")
    assert TranscriptDOCXHandler().can_handle(fake) is False


def test_can_handle_returns_false_for_non_docx_suffix(tmp_path: Path) -> None:
    """Non-.docx files are rejected on suffix before any content sniff."""
    h = TranscriptDOCXHandler()
    txt = tmp_path / "x.txt"
    txt.write_text("Alice: hi\nBob: hello\nAlice: how are you", encoding="utf-8")
    assert h.can_handle(txt) is False
    pdf = tmp_path / "x.pdf"
    pdf.write_bytes(b"%PDF-1.4 fake")
    assert h.can_handle(pdf) is False


def test_can_handle_returns_false_for_missing_file(tmp_path: Path) -> None:
    """Non-existent .docx paths are rejected before opening."""
    missing = tmp_path / "does-not-exist.docx"
    assert TranscriptDOCXHandler().can_handle(missing) is False


def test_can_handle_returns_false_for_str_input(tmp_path: Path) -> None:
    """can_handle is Path-only; raw strings are rejected (URLs etc. are not docx)."""
    assert TranscriptDOCXHandler().can_handle("not-a-path.docx") is False


def test_can_handle_rejects_prose_with_mid_sentence_colons(tmp_path: Path) -> None:
    """Prose paragraphs with mid-sentence colons must NOT false-match.

    The speaker regex is capped at 3 tokens before the colon, so sentences
    like "This is something I learned today: ..." (5+ tokens before the
    colon) don't satisfy the sniff. This is the primary false-positive
    risk for narrowing TranscriptDOCXHandler.can_handle.
    """
    doc = Document()
    doc.add_paragraph("Here is a thought I had today: the world is round.")
    doc.add_paragraph("Another idea worth considering carefully: we need a plan.")
    doc.add_paragraph("Something else I noticed this morning: the weather changed.")
    doc.add_paragraph("One more observation about the meeting: it ran long.")
    doc.add_paragraph("Final reflection from this whole exercise: lessons learned.")
    path = tmp_path / "prose-with-colons.docx"
    doc.save(str(path))
    assert TranscriptDOCXHandler().can_handle(path) is False


def test_can_handle_claims_qa_interview_format(tmp_path: Path) -> None:
    """Interview-style Q:/A: docs are transcript-adjacent and welcome to claim.

    Pinning the intentional design choice that short label-like prefixes
    (Q:, A:, Note:, TODO:) count as speaker-shape. If this proves noisy
    in real-world docs, tighten the regex.
    """
    doc = Document()
    doc.add_paragraph("Q: What is your background?")
    doc.add_paragraph("A: I have ten years of experience in product.")
    doc.add_paragraph("Q: How did you get started?")
    doc.add_paragraph("A: I joined a startup right out of school.")
    path = tmp_path / "qa.docx"
    doc.save(str(path))
    assert TranscriptDOCXHandler().can_handle(path) is True
