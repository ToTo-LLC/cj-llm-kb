"""Unit tests for DocxHandler (Plan 24 T1).

Fixtures are constructed at runtime via python-docx so the repo stays free
of binary .docx files. The single existing binary fixture
(`tests/ingest/fixtures/notes.docx`) is owned by the transcript handler.
"""

from __future__ import annotations

import struct
import zlib
from pathlib import Path

import pytest
from brain_core.ingest.dispatcher import dispatch
from brain_core.ingest.handlers.base import HandlerError
from brain_core.ingest.handlers.docx import DocxHandler
from brain_core.ingest.handlers.transcript_docx import TranscriptDOCXHandler
from brain_core.ingest.types import SourceType
from docx import Document

# ---------------------------------------------------------------------------
# Test helpers
# ---------------------------------------------------------------------------


def _make_tiny_png() -> bytes:
    """Build a valid 1x1 white RGB PNG using stdlib zlib + struct.

    We avoid PIL / fixture binary files so the test is hermetic.
    """
    sig = b"\x89PNG\r\n\x1a\n"
    # IHDR: width=1, height=1, bit_depth=8, color_type=2 (RGB),
    # compression=0, filter=0, interlace=0
    ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    ihdr = (
        struct.pack(">I", 13)
        + b"IHDR"
        + ihdr_data
        + struct.pack(">I", zlib.crc32(b"IHDR" + ihdr_data))
    )
    # IDAT: filter byte 0 + RGB white
    raw = b"\x00\xff\xff\xff"
    compressed = zlib.compress(raw)
    idat = (
        struct.pack(">I", len(compressed))
        + b"IDAT"
        + compressed
        + struct.pack(">I", zlib.crc32(b"IDAT" + compressed))
    )
    iend = (
        struct.pack(">I", 0)
        + b"IEND"
        + struct.pack(">I", zlib.crc32(b"IEND"))
    )
    return sig + ihdr + idat + iend


def _build_plain_docx(tmp_path: Path, name: str = "doc.docx") -> Path:
    doc = Document()
    doc.add_paragraph("First paragraph here.")
    doc.add_paragraph("Second paragraph.")
    doc.add_paragraph("Third and final paragraph.")
    path = tmp_path / name
    doc.save(str(path))
    return path


def _build_headings_docx(tmp_path: Path, name: str = "headings.docx") -> Path:
    doc = Document()
    doc.add_heading("Strategy 2026", level=1)
    doc.add_heading("Quarterly themes", level=2)
    doc.add_heading("Q1 focus", level=3)
    doc.add_paragraph("Body content under H3.")
    path = tmp_path / name
    doc.save(str(path))
    return path


def _build_table_docx(tmp_path: Path, name: str = "table.docx") -> Path:
    doc = Document()
    doc.add_paragraph("Before table.")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "Metric"
    tbl.cell(0, 1).text = "Value"
    tbl.cell(1, 0).text = "Revenue"
    tbl.cell(1, 1).text = "100"
    doc.add_paragraph("After table.")
    path = tmp_path / name
    doc.save(str(path))
    return path


def _build_image_docx(tmp_path: Path, name: str = "with-image.docx") -> Path:
    doc = Document()
    doc.add_paragraph("Some text above the image.")
    img_path = tmp_path / "_pix.png"
    img_path.write_bytes(_make_tiny_png())
    doc.add_picture(str(img_path))
    path = tmp_path / name
    doc.save(str(path))
    return path


def _build_empty_docx(tmp_path: Path, name: str = "empty.docx") -> Path:
    doc = Document()
    path = tmp_path / name
    doc.save(str(path))
    return path


def _build_no_title_docx(tmp_path: Path, name: str = "no-title.docx") -> Path:
    """Document with no Heading 1 / Title — first content is a plain paragraph."""
    doc = Document()
    doc.add_paragraph("Just a plain paragraph.")
    doc.add_paragraph("Another plain paragraph.")
    path = tmp_path / name
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Dispatcher routing — Plan 24 D3
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_dispatcher_routes_transcript_docx_first(tmp_path: Path) -> None:
    """Per D3: TranscriptDOCXHandler claims FIRST. DocxHandler is fall-through.

    Today TranscriptDOCXHandler.can_handle claims ANY .docx by suffix, so a
    transcript-stem-named file still goes to it. This test pins that
    invariant so future narrowing of TranscriptDOCXHandler.can_handle (if
    it adopts a stem convention) still keeps the same example file routed
    to the transcript handler.
    """
    path = _build_plain_docx(tmp_path, name="interview-2024-05.docx")
    handler = await dispatch(path)
    assert isinstance(handler, TranscriptDOCXHandler)


@pytest.mark.asyncio
async def test_dispatcher_routes_generic_docx_to_docx_handler(tmp_path: Path) -> None:
    """When TranscriptDOCXHandler is bypassed, generic .docx flows to DocxHandler.

    TranscriptDOCXHandler currently claims all .docx, so to verify the
    fall-through wiring we pass a handler list that omits the transcript
    handler. This is the same pattern the dispatcher will exhibit once
    TranscriptDOCXHandler is narrowed to transcript-stem only.
    """
    path = _build_plain_docx(tmp_path, name="Q4-strategy.docx")
    handler = await dispatch(
        path,
        handlers=[DocxHandler()],
    )
    assert isinstance(handler, DocxHandler)


def test_docx_handler_registered_after_transcript_docx() -> None:
    """Order check on the actual default handler chain.

    Pins the dispatcher's `_default_handlers()` ordering: TranscriptDOCXHandler
    must precede DocxHandler. Required by Plan 24 D3 — D3 collapses if
    DocxHandler ever moves up.
    """
    from brain_core.ingest.dispatcher import _default_handlers

    chain = _default_handlers()
    transcript_idx = next(
        i for i, h in enumerate(chain) if isinstance(h, TranscriptDOCXHandler)
    )
    docx_idx = next(i for i, h in enumerate(chain) if isinstance(h, DocxHandler))
    assert transcript_idx < docx_idx


def test_docx_handler_rejects_non_path() -> None:
    """can_handle is a pure routing check; rejects raw strings + non-.docx paths."""
    h = DocxHandler()
    assert h.can_handle("https://example.com/foo.docx") is False
    assert h.can_handle(Path("/tmp/foo.pdf")) is False
    assert h.can_handle(Path("/tmp/foo.DOCX")) is True  # case-insensitive


# ---------------------------------------------------------------------------
# Extraction
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_extract_paragraphs_preserved(tmp_path: Path) -> None:
    path = _build_plain_docx(tmp_path)
    archive = tmp_path / "archive"
    h = DocxHandler()
    es = await h.extract(path, archive_root=archive)
    assert es.source_type is SourceType.DOCX
    assert "First paragraph here." in es.body_text
    assert "Second paragraph." in es.body_text
    assert "Third and final paragraph." in es.body_text
    assert es.archive_path.exists()
    assert es.archive_path.parent == archive


@pytest.mark.asyncio
async def test_extract_headings_to_markdown(tmp_path: Path) -> None:
    path = _build_headings_docx(tmp_path)
    es = await DocxHandler().extract(path, archive_root=tmp_path / "archive")
    # Heading 1 → "# ", Heading 2 → "## ", Heading 3 → "### "
    assert "# Strategy 2026" in es.body_text
    assert "## Quarterly themes" in es.body_text
    assert "### Q1 focus" in es.body_text
    # Title fallback chain picks the H1 text
    assert es.title == "Strategy 2026"


@pytest.mark.asyncio
async def test_extract_tables_to_gfm_markdown(tmp_path: Path) -> None:
    path = _build_table_docx(tmp_path)
    es = await DocxHandler().extract(path, archive_root=tmp_path / "archive")
    # GFM separator line confirms valid markdown table emit
    assert "| Metric | Value |" in es.body_text
    assert "| --- | --- |" in es.body_text
    assert "| Revenue | 100 |" in es.body_text
    # Interleave preserved: table appears between the two paragraphs
    before_idx = es.body_text.find("Before table.")
    table_idx = es.body_text.find("| Metric")
    after_idx = es.body_text.find("After table.")
    assert before_idx < table_idx < after_idx


@pytest.mark.asyncio
async def test_extract_collects_inline_images(tmp_path: Path) -> None:
    path = _build_image_docx(tmp_path)
    es = await DocxHandler().extract(path, archive_root=tmp_path / "archive")
    images = es.extras.get("images", [])
    assert len(images) == 1
    img = images[0]
    # T4's OCR pass expects this exact shape.
    assert isinstance(img["blob"], bytes)
    assert len(img["blob"]) > 0
    assert img["content_type"] == "image/png"
    assert img["index"] == 0


@pytest.mark.asyncio
async def test_extract_empty_docx(tmp_path: Path) -> None:
    path = _build_empty_docx(tmp_path)
    es = await DocxHandler().extract(path, archive_root=tmp_path / "archive")
    assert es.body_text == ""
    assert es.extras.get("images", []) == []
    # Title falls back to the filename stem when the doc has no content.
    assert es.title == "empty"
    assert es.source_type is SourceType.DOCX


@pytest.mark.asyncio
async def test_extract_uses_filename_stem_when_no_title(tmp_path: Path) -> None:
    """No H1 / Title → first non-empty paragraph wins the title slot."""
    path = _build_no_title_docx(tmp_path, name="memo-2026.docx")
    es = await DocxHandler().extract(path, archive_root=tmp_path / "archive")
    # First non-empty paragraph fills the title slot (per plan §T1.3).
    assert es.title == "Just a plain paragraph."


@pytest.mark.asyncio
async def test_extract_raises_handler_error_on_corrupt_docx(tmp_path: Path) -> None:
    """Garbage with .docx extension must raise HandlerError, not crash."""
    fake = tmp_path / "fake.docx"
    fake.write_bytes(b"not a real docx file")
    with pytest.raises(HandlerError):
        await DocxHandler().extract(fake, archive_root=tmp_path / "archive")


@pytest.mark.asyncio
async def test_extract_raises_when_file_missing(tmp_path: Path) -> None:
    missing = tmp_path / "missing.docx"
    with pytest.raises(HandlerError):
        await DocxHandler().extract(missing, archive_root=tmp_path / "archive")
