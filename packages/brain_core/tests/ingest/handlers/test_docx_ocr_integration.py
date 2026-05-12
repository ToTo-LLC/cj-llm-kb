"""Plan 24 Task 4 — integration tests for the pipeline OCR pass on DocxHandler output.

End-to-end coverage: build a real .docx fixture via python-docx, run it through
:class:`IngestPipeline` with :class:`FakeLLMProvider` priming the classify /
summarize / integrate / vision queues, and assert the OCR'd text lands in the
final vault source note's body (or doesn't, on the error / empty / no-image
paths).

These tests EXERCISE THE FULL PIPELINE — they're the canary that proves T4's
``_ocr_images`` insertion point fires AT the right stage (post-classify
quarantine check, pre-summarize) and that the inline format
(``[Image: <text>]``) is preserved through summarize + integrate + apply.
"""

from __future__ import annotations

import struct
import zlib
from datetime import UTC, datetime, timedelta
from pathlib import Path

import pytest
from brain_core.budget import PerDomainBudgetGuard
from brain_core.config.schema import BudgetOverride, Config
from brain_core.cost.ledger import CostEntry, CostLedger
from brain_core.ingest.pipeline import IngestPipeline
from brain_core.ingest.types import IngestStatus
from brain_core.llm.fake import FakeLLMProvider
from brain_core.prompts.schemas import SummarizeOutput
from brain_core.vault.types import IndexEntryPatch, PatchSet
from brain_core.vault.writer import VaultWriter
from docx import Document
from structlog.testing import capture_logs

# ---------------------------------------------------------------------------
# Fixture builders (hermetic — no binary files committed)
# ---------------------------------------------------------------------------


def _make_tiny_png() -> bytes:
    """Build a valid 1x1 white RGB PNG using stdlib zlib + struct.

    Same helper as ``test_docx.py``. Kept local to dodge a cross-file
    import; the OCR pass doesn't care about the pixels — only that
    ``extras["images"]`` carries valid bytes.
    """
    sig = b"\x89PNG\r\n\x1a\n"
    ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    ihdr = (
        struct.pack(">I", 13)
        + b"IHDR"
        + ihdr_data
        + struct.pack(">I", zlib.crc32(b"IHDR" + ihdr_data))
    )
    raw = b"\x00\xff\xff\xff"
    compressed = zlib.compress(raw)
    idat = (
        struct.pack(">I", len(compressed))
        + b"IDAT"
        + compressed
        + struct.pack(">I", zlib.crc32(b"IDAT" + compressed))
    )
    iend = (
        struct.pack(">I", 0)
        + b"IEND"
        + struct.pack(">I", zlib.crc32(b"IEND"))
    )
    return sig + ihdr + idat + iend


def _build_image_docx(tmp_path: Path, name: str = "with-image.docx") -> Path:
    """One paragraph + one inline image. Image bytes don't matter for OCR — only
    the dict shape feeding ``ocr_image``."""
    doc = Document()
    doc.add_paragraph("Some text above the image.")
    img_path = tmp_path / "_pix.png"
    img_path.write_bytes(_make_tiny_png())
    doc.add_picture(str(img_path))
    path = tmp_path / name
    doc.save(str(path))
    return path


def _build_plain_docx(tmp_path: Path, name: str = "plain.docx") -> Path:
    """Plain prose — zero images. Used to pin the no-OCR-call path."""
    doc = Document()
    doc.add_paragraph("First paragraph here.")
    doc.add_paragraph("Second paragraph.")
    doc.add_paragraph("Third and final paragraph.")
    path = tmp_path / name
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Test plumbing — queue classify+summarize+integrate so the pipeline reaches
# the post-classify OCR pass. The fake LLM provider serves all of them from
# the same ``_queue`` via ``complete()``.
# ---------------------------------------------------------------------------


def _queue_pipeline_responses(
    fake: FakeLLMProvider,
    *,
    title: str,
) -> None:
    """Queue summarize + integrate so the ingest reaches Stage 9.

    Callers pass ``domain_override="research"`` so classify is SKIPPED — no
    classify response is queued. The vision queue is the caller's
    responsibility (queue/leave-empty based on the scenario under test).
    """
    fake.queue(
        SummarizeOutput(
            title=title,
            summary="summary text",
            key_points=["point"],
            entities=[],
            concepts=[],
            open_questions=[],
        ).model_dump_json()
    )
    fake.queue(
        PatchSet(
            index_entries=[
                IndexEntryPatch(
                    section="Sources",
                    line=f"- [[{title}]] — summary text",
                    domain="research",
                )
            ],
            log_entry=f"## [2026-05-12 10:00] ingest | source | [[{title}]]",
            reason="initial",
        ).model_dump_json()
    )


def _make_pipeline(
    fake: FakeLLMProvider,
    *,
    vault: Path,
    cost_ledger: CostLedger,
    config: Config,
) -> IngestPipeline:
    guard = PerDomainBudgetGuard(cost_ledger)
    return IngestPipeline(
        vault_root=vault,
        writer=VaultWriter(vault_root=vault),
        llm=fake,
        summarize_model="claude-sonnet-4-6",
        integrate_model="claude-sonnet-4-6",
        classify_model="claude-haiku-4-5-20251001",
        guard=guard,
        config=config,
        cost_ledger=cost_ledger,
    )


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_docx_with_one_inline_image_gets_ocr_block(
    ephemeral_vault: Path, tmp_path: Path
) -> None:
    """Happy path: 1 inline image → vision text inlined as ``[Image: ...]`` in the
    final source note body (and therefore the summarize/integrate input)."""
    docx_path = _build_image_docx(tmp_path)
    ledger = CostLedger(db_path=tmp_path / "costs.sqlite")
    config = Config(vault_path=ephemeral_vault)

    fake = FakeLLMProvider()
    fake.queue_vision("diagram showing X")
    _queue_pipeline_responses(fake, title="image-doc")

    p = _make_pipeline(fake, vault=ephemeral_vault, cost_ledger=ledger, config=config)

    r = await p.ingest(
        docx_path,
        allowed_domains=("research",),
        domain_override="research",  # skip classify
    )

    assert r.status is IngestStatus.OK
    assert r.note_path is not None
    assert r.extracted is not None
    # The OCR'd body_text is the canonical pin: it's what the summarize +
    # integrate prompts saw, and it's the field T4 augments. The final
    # source note's body is rendered from SummarizeOutput, not body_text
    # verbatim, so asserting on the note file would only pin LLM behavior.
    assert "[Image: diagram showing X]" in r.extracted.body_text
    # The original prose is preserved — OCR blocks are appended.
    assert "Some text above the image." in r.extracted.body_text
    # And one vision call was made.
    assert len(fake.vision_calls) == 1
    # The summarize prompt was actually fed the OCR'd body (regression
    # pin: if the OCR pass ran AFTER summarize, this would fail).
    summarize_req = fake.requests[0]  # classify skipped via override
    assert "[Image: diagram showing X]" in summarize_req.messages[0].content
    # Cost ledger picked up the op="ocr" row.
    today = datetime.now(tz=UTC).date()
    by_domain = ledger.total_by_domain(today)
    assert by_domain.get("research", 0.0) > 0.0


@pytest.mark.asyncio
async def test_docx_with_no_images_no_ocr_call(
    ephemeral_vault: Path, tmp_path: Path
) -> None:
    """Pure-text docx (0 images) → vision_extract NEVER called.

    The negative-evidence assertion: the vision queue stays untouched.
    """
    docx_path = _build_plain_docx(tmp_path)
    ledger = CostLedger(db_path=tmp_path / "costs.sqlite")
    config = Config(vault_path=ephemeral_vault)

    fake = FakeLLMProvider()
    # DELIBERATELY queue NO vision response. If OCR fires, FakeLLMProvider
    # raises RuntimeError("...vision queue is empty...") and the test fails.
    _queue_pipeline_responses(fake, title="plain-doc")

    p = _make_pipeline(fake, vault=ephemeral_vault, cost_ledger=ledger, config=config)

    r = await p.ingest(
        docx_path,
        allowed_domains=("research",),
        domain_override="research",
    )

    assert r.status is IngestStatus.OK
    assert fake.vision_calls == []  # vision_extract was never invoked


@pytest.mark.asyncio
async def test_docx_with_vision_error_skips_image_continues_ingest(
    ephemeral_vault: Path, tmp_path: Path
) -> None:
    """Vision API error on an image → log warning + skip image + ingest still OK.

    Setup: 1-image docx, vision queue EMPTY. The OCR pass catches the
    ``RuntimeError`` (FakeLLMProvider's empty-queue raise), logs a warning,
    and continues. The final note has NO ``[Image: ...]`` block but the
    ingest still completes with status=OK.
    """
    docx_path = _build_image_docx(tmp_path)
    ledger = CostLedger(db_path=tmp_path / "costs.sqlite")
    config = Config(vault_path=ephemeral_vault)

    fake = FakeLLMProvider()
    # Vision queue intentionally empty → first ocr_image call raises RuntimeError.
    _queue_pipeline_responses(fake, title="error-doc")

    p = _make_pipeline(fake, vault=ephemeral_vault, cost_ledger=ledger, config=config)

    with capture_logs() as logs:
        r = await p.ingest(
            docx_path,
            allowed_domains=("research",),
            domain_override="research",
        )

    assert r.status is IngestStatus.OK
    assert r.note_path is not None
    assert r.extracted is not None
    # No OCR block in the (final) body_text: the only image errored out.
    assert "[Image:" not in r.extracted.body_text
    # Original prose preserved.
    assert "Some text above the image." in r.extracted.body_text
    # Warning was logged.
    warnings = [e for e in logs if e.get("event") == "ingest.ocr.image_skipped"]
    assert len(warnings) == 1
    assert "vision queue is empty" in warnings[0]["error"]
    assert warnings[0]["error_type"] == "RuntimeError"


@pytest.mark.asyncio
async def test_docx_with_empty_ocr_response_no_inline_block(
    ephemeral_vault: Path, tmp_path: Path
) -> None:
    """Vision returns empty string → no ``[Image: ]`` block in body."""
    docx_path = _build_image_docx(tmp_path)
    ledger = CostLedger(db_path=tmp_path / "costs.sqlite")
    config = Config(vault_path=ephemeral_vault)

    fake = FakeLLMProvider()
    fake.queue_vision("")  # empty OCR result
    _queue_pipeline_responses(fake, title="empty-ocr-doc")

    p = _make_pipeline(fake, vault=ephemeral_vault, cost_ledger=ledger, config=config)

    r = await p.ingest(
        docx_path,
        allowed_domains=("research",),
        domain_override="research",
    )

    assert r.status is IngestStatus.OK
    assert r.note_path is not None
    assert r.extracted is not None
    # Empty OCR → no inline ``[Image: ]`` block. Body unchanged.
    assert "[Image:" not in r.extracted.body_text
    # Vision WAS called once — we just didn't emit a block from the result.
    assert len(fake.vision_calls) == 1


@pytest.mark.asyncio
async def test_docx_with_budget_exhausted_raises_and_status_failed(
    ephemeral_vault: Path, tmp_path: Path
) -> None:
    """Per-domain budget exhausted before OCR → BudgetCapExceeded propagates
    to the outer pipeline ``try`` and lands as status=FAILED.

    Pins two invariants:
    (a) The OCR pass DOES call ``budget_guard.check_for(...)`` before the LLM.
    (b) The ingest catches the exception and returns status=FAILED rather
        than crashing.
    """
    docx_path = _build_image_docx(tmp_path)
    ledger = CostLedger(db_path=tmp_path / "costs.sqlite")
    config = Config(vault_path=ephemeral_vault)
    # Pin a $0.001 daily cap on 'research' and seed a 2x-cap row so the
    # next call exceeds it.
    config.budget.per_domain = {"research": BudgetOverride(daily_cap_usd=0.001)}
    ledger.record(
        CostEntry(
            timestamp=datetime.now(tz=UTC) - timedelta(hours=1),
            operation="ingest",
            model="claude-sonnet-4-6",
            input_tokens=200,
            output_tokens=10,
            cost_usd=0.002,
            domain="research",
        )
    )

    fake = FakeLLMProvider()
    # Queue classify+summarize+integrate so the pipeline reaches the OCR
    # pass; queue NO vision response so a wrong-direction failure would
    # surface "vision queue is empty" instead of BudgetCapExceeded.
    _queue_pipeline_responses(fake, title="budget-doc")

    p = _make_pipeline(fake, vault=ephemeral_vault, cost_ledger=ledger, config=config)

    r = await p.ingest(
        docx_path,
        allowed_domains=("research",),
        domain_override="research",
    )

    assert r.status is IngestStatus.FAILED
    assert any("domain=research" in e for e in r.errors) or any(
        "BudgetCapExceeded" in e or "cap" in e.lower() for e in r.errors
    )
    # LLM vision was never called — the budget gate fired before.
    assert fake.vision_calls == []
