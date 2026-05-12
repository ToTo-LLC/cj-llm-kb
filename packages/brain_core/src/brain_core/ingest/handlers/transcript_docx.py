"""DOCX transcript handler — reads paragraphs with python-docx.

Plan 24 T1.5: ``can_handle`` content-sniffs the document to claim only
transcript-shaped ``.docx`` files (speaker labels or timestamps in the
first paragraphs). Generic Word documents fall through to DocxHandler
per Plan 24 D3. Corrupt or unreadable files return False rather than
raising — the dispatcher then tries the next handler (DocxHandler), which
emits a clean HandlerError at ``extract`` time if the file is also bad.
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document

from brain_core.ingest.handlers.base import HandlerError
from brain_core.ingest.types import ExtractedSource, SourceType

# Speaker labels at start of paragraph. Matches:
#   "Alice:"  "John Doe:"  "Speaker 1:"  "Dr. Smith:"  "Mary Jane Watson:"
#   "Q:"  "A:"  "Note:"  "TODO:"  — interview / outline formats are
#   transcript-adjacent and welcome to claim.
# Rejects:
#   "This is something I learned today: ..." — prose with mid-sentence
#   colon (more than 3 tokens before the colon).
# Capped at 3 space-separated tokens before the colon. Word boundary
# after the colon (whitespace or EOL) keeps prose with mid-sentence
# colons from false-matching.
_SPEAKER_PATTERN = re.compile(r"^[A-Z][\w.'\-]*(?:\s[\w.'\-]+){0,2}:(?:\s|$)")

# All-caps bracketed names (Zoom, some Otter exports): "[CHRIS JOHNSON]"
_BRACKET_SPEAKER_PATTERN = re.compile(r"^\[[A-Z][A-Z0-9 .'\-]+\]")

# Timestamps at start of paragraph. Matches:
#   "01:23"  "12:34:56"  "[01:23]"  "[12:34:56]"  "(01:23)"
_TIMESTAMP_PATTERN = re.compile(r"^[\[\(]?\d{1,2}:\d{2}(?::\d{2})?[\]\)]?")

# ">>" speaker markers used by auto-transcription services.
_ARROW_SPEAKER_PATTERN = re.compile(r"^>>\s")

_TRANSCRIPT_MIN_MATCHES = 3  # need 3+ transcript-shaped paragraphs to claim
_TRANSCRIPT_SNIFF_WINDOW = 10  # search first N non-empty paragraphs


def _looks_like_transcript_line(text: str) -> bool:
    """True iff the line starts with a speaker label, timestamp, or arrow marker."""
    return bool(
        _SPEAKER_PATTERN.match(text)
        or _BRACKET_SPEAKER_PATTERN.match(text)
        or _TIMESTAMP_PATTERN.match(text)
        or _ARROW_SPEAKER_PATTERN.match(text)
    )


class TranscriptDOCXHandler:
    source_type: SourceType = SourceType.TRANSCRIPT

    def can_handle(self, spec: str | Path) -> bool:
        """Content-sniff: only claim ``.docx`` files that look like transcripts.

        Requires ``_TRANSCRIPT_MIN_MATCHES`` (3) speaker-label / timestamp /
        arrow-marker hits in the first ``_TRANSCRIPT_SNIFF_WINDOW`` (10)
        non-empty paragraphs. Generic Word documents (no speakers /
        timestamps) fall through to ``DocxHandler`` per Plan 24 D3.
        Corrupt or unreadable ``.docx`` files return ``False`` — the
        dispatcher continues to the next handler, which surfaces a clean
        ``HandlerError`` at extract time.

        Pure path / file-shape checks happen before opening the doc so we
        only pay the python-docx open cost on actual ``.docx`` files.
        """
        if not isinstance(spec, Path):
            return False
        if spec.suffix.lower() != ".docx":
            return False
        if not spec.is_file():
            return False
        try:
            doc = Document(str(spec))
        except Exception:
            # Corrupt or non-docx ZIP — don't claim. DocxHandler will
            # surface the actionable HandlerError at extract time.
            return False
        nonempty: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            nonempty.append(text)
            if len(nonempty) >= _TRANSCRIPT_SNIFF_WINDOW:
                break
        if len(nonempty) < _TRANSCRIPT_MIN_MATCHES:
            return False
        matches = sum(1 for line in nonempty if _looks_like_transcript_line(line))
        return matches >= _TRANSCRIPT_MIN_MATCHES

    async def extract(self, spec: str | Path, *, archive_root: Path) -> ExtractedSource:
        if not isinstance(spec, Path) or not spec.exists():
            raise HandlerError(f"transcript_docx cannot read {spec!r}")
        try:
            doc = Document(spec)  # type: ignore[arg-type]  # python-docx stubs say str | IO[bytes], but Path works at runtime
        except Exception as exc:
            raise HandlerError(
                f"Could not open DOCX {spec.name!r}: {exc}. "
                "The file may be corrupt or not a Word document."
            ) from exc
        body = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        archive_root.mkdir(parents=True, exist_ok=True)
        archive_path = archive_root / spec.name
        shutil.copy2(spec, archive_path)
        return ExtractedSource(
            title=spec.stem,
            author=None,
            published=None,
            source_url=None,
            source_type=SourceType.TRANSCRIPT,
            body_text=body,
            archive_path=archive_path,
        )
