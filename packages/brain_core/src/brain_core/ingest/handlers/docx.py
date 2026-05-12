"""DocxHandler — generic Word (.docx) extraction.

Plan 24 T1. Generic non-transcript Microsoft Word document handler.

Routing contract (Plan 24 D3): `TranscriptDOCXHandler` is registered FIRST in
the dispatcher and claims all `.docx` files. `DocxHandler` is the fall-through
— it claims any `.docx` by suffix, but the dispatcher will never reach it for
files the transcript handler has already claimed. If `TranscriptDOCXHandler`
is later narrowed to a stem convention, generic `.docx` will naturally flow
to this handler with no code change here.

Extraction shape (per plan §T1.3):
- Paragraphs in document order; Heading 1..6 / Title styles emit markdown
  `#`..`######`. Other styles emit as plain paragraphs.
- Tables emit as GFM markdown tables, interleaved with paragraphs in
  document order via `doc.element.body` iteration (otherwise tables would
  bunch at the end).
- Inline images collected into `extras["images"]` as
  `{"blob": bytes, "content_type": str, "index": int}` dicts. The
  pipeline OCR pass (Plan 24 T4) consumes this shape.
- `title`: first Heading 1 / Title paragraph text; falls back to the first
  non-empty paragraph; falls back to the file stem if the document is empty.
- `author` / `published`: from `doc.core_properties` (author may be empty
  string per python-docx; we coerce to None for unset).
- `source_url`: None (docx is local).
- `source_type`: `SourceType.DOCX`.
- `archive_path`: original file copied into `archive_root` (mirrors
  `TranscriptDOCXHandler`).
"""

from __future__ import annotations

import shutil
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from brain_core.ingest.handlers.base import HandlerError
from brain_core.ingest.types import ExtractedSource, SourceType

# Paragraph styles we map to a markdown heading. Anything else renders as a
# plain paragraph. python-docx exposes `paragraph.style.name`; the default
# heading styles in Word are literal "Heading 1".."Heading 9" plus "Title".
_HEADING_STYLES: dict[str, int] = {
    "Title": 1,
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
    "Heading 5": 5,
    "Heading 6": 6,
}


class DocxHandler:
    """Generic Word document handler. Fall-through after TranscriptDOCXHandler."""

    source_type: SourceType = SourceType.DOCX

    def can_handle(self, spec: str | Path) -> bool:
        """Claim any `.docx` file. Pure routing check — no I/O."""
        return isinstance(spec, Path) and spec.suffix.lower() == ".docx"

    async def extract(self, spec: str | Path, *, archive_root: Path) -> ExtractedSource:
        """Open the .docx, walk body in document order, collect images."""
        if not isinstance(spec, Path) or not spec.exists():
            raise HandlerError(f"docx handler cannot read {spec!r}")
        try:
            doc: DocumentType = Document(str(spec))
        except Exception as exc:
            raise HandlerError(
                f"Could not open DOCX {spec.name!r}: {exc}. "
                "The file may be corrupt or not a Word document."
            ) from exc

        body_parts: list[str] = []
        title: str | None = None
        first_nonempty_paragraph: str | None = None

        # Walk the body in document order so paragraphs + tables interleave
        # correctly. `doc.paragraphs` and `doc.tables` are separate views and
        # would lose interleave context if joined naively.
        for block in _iter_body_blocks(doc):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                style_name = block.style.name if block.style is not None else ""
                heading_level = _HEADING_STYLES.get(style_name)
                if text:
                    if heading_level is not None:
                        if title is None:
                            title = text
                        body_parts.append(f"{'#' * heading_level} {text}")
                    else:
                        if first_nonempty_paragraph is None:
                            first_nonempty_paragraph = text
                        body_parts.append(text)
                # Empty paragraphs (including the empty paragraph
                # python-docx injects after inline pictures) are skipped.
            elif isinstance(block, Table):
                rendered = _table_to_gfm_markdown(block)
                if rendered:
                    body_parts.append(rendered)

        body_text = "\n\n".join(body_parts)

        # Collect inline images for the T4 OCR pass.
        images: list[dict[str, Any]] = []
        for idx, shape in enumerate(doc.inline_shapes):
            blob_entry = _extract_inline_image(doc, shape, idx)
            if blob_entry is not None:
                images.append(blob_entry)

        # Title fallback chain: Heading 1/Title -> first non-empty para -> stem.
        if title is None:
            title = first_nonempty_paragraph or spec.stem

        # Core props: author / created. python-docx returns "" for unset author
        # and None for unset created.
        author_raw = doc.core_properties.author
        author: str | None = author_raw.strip() if author_raw else None
        if author == "":
            author = None
        created_raw = doc.core_properties.created
        published: date | None = created_raw.date() if created_raw is not None else None

        archive_root.mkdir(parents=True, exist_ok=True)
        archive_path = archive_root / spec.name
        shutil.copy2(spec, archive_path)

        extras: dict[str, Any] = {"images": images}

        return ExtractedSource(
            title=title,
            author=author,
            published=published,
            source_url=None,
            source_type=SourceType.DOCX,
            body_text=body_text,
            archive_path=archive_path,
            extras=extras,
        )


def _iter_body_blocks(doc: DocumentType) -> list[Paragraph | Table]:
    """Yield paragraphs and tables in document order.

    `doc.paragraphs` + `doc.tables` are separate views; iterating
    `doc.element.body.iterchildren()` preserves authoring order, which we
    need so a table inserted between two paragraphs renders where the
    author put it (not bunched at the end).

    The Paragraph/Table parent is the Document itself — `Paragraph.part`
    walks `_parent.part`, and only the Document (not the raw `<w:body>`
    XML element) exposes that attribute.
    """
    out: list[Paragraph | Table] = []
    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")
    for child in doc.element.body.iterchildren():
        if child.tag == p_tag:
            out.append(Paragraph(child, doc))  # type: ignore[arg-type]
        elif child.tag == tbl_tag:
            out.append(Table(child, doc))  # type: ignore[arg-type]
        # other elements (sectPr, etc.) intentionally skipped
    return out


def _table_to_gfm_markdown(table: Table) -> str:
    """Render a python-docx Table as a GFM markdown table.

    Header row: row 0 (Word doesn't enforce a header convention; row 0 is
    conventional). For a single-row table we still emit the GFM separator
    so the result is a valid table.

    Cells with embedded pipes are escaped to keep the table well-formed.
    Newlines inside a cell are flattened to a space.
    """
    if len(table.rows) == 0:
        return ""
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [_escape_cell(cell.text) for cell in row.cells]
        rows.append(cells)
    col_count = max(len(r) for r in rows)
    # Pad rows that have fewer cells (rare but possible with merged cells).
    for r in rows:
        while len(r) < col_count:
            r.append("")
    header = "| " + " | ".join(rows[0]) + " |"
    separator = "| " + " | ".join("---" for _ in range(col_count)) + " |"
    body = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:])
    if body:
        return "\n".join([header, separator, body])
    return "\n".join([header, separator])


def _escape_cell(text: str) -> str:
    """Escape GFM table cell contents (pipes + newlines)."""
    return text.replace("\n", " ").replace("|", "\\|").strip()


def _extract_inline_image(
    doc: DocumentType,
    shape: Any,
    index: int,
) -> dict[str, Any] | None:
    """Pull the image blob + content_type for a single InlineShape.

    python-docx's image API is roundabout: the InlineShape's underlying
    `<a:blip>` element holds an embed-rId that resolves through
    `doc.part.related_parts` to the ImagePart, which exposes `.blob` and
    `.content_type`. Non-PICTURE shapes (charts, smart art, linked
    pictures) return None so the caller can skip them.
    """
    try:
        from docx.enum.shape import WD_INLINE_SHAPE

        if shape.type != WD_INLINE_SHAPE.PICTURE:
            return None
        blip = shape._inline.graphic.graphicData.pic.blipFill.blip
        r_id = blip.embed
        if not r_id:
            return None
        image_part = doc.part.related_parts[r_id]
        return {
            "blob": image_part.blob,
            "content_type": image_part.content_type,
            "index": index,
        }
    except (AttributeError, KeyError):
        # Defensive: malformed embedded shape — skip rather than failing
        # the whole extract.
        return None
